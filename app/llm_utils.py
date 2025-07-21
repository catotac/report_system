from app.models import DocumentRequest, DocumentResult, SectionOutput, EvaluationResult
from app.prompts import load_template, load_group_template
import openai
import re
from docx import Document
from fastapi.responses import FileResponse
import tempfile

openai.api_key = ""##

def call_openai(prompt: str) -> str:
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}]
    )
    return response['choices'][0]['message']['content']

def evaluate_generation(text: str) -> EvaluationResult:
    eval_prompt = f"Evaluate the following for groundedness, completeness, coherence from 0 to 1:\n{text}"
    eval_response = call_openai(eval_prompt)
    def extract_first_float(line):
        match = re.search(r"[-+]?[0-9]*\.?[0-9]+", line)
        return float(match.group()) if match else 0.0
    lines = eval_response.split("\n")
    scores = []
    for line in lines:
        if len(scores) >= 3:
            break
        scores.append(extract_first_float(line))
    while len(scores) < 3:
        scores.append(0.0)
    return EvaluationResult(groundedness=scores[0], completeness=scores[1], coherence=scores[2])

def self_reflect_and_improve(text: str, evaluation: EvaluationResult, section: str, subsection: str, context: str = "", group_id: str = "default") -> str:
    """Use evaluation scores to improve the previous generation through self-reflection."""
    prompt_template = load_group_template("report", group_id, section, subsection)
    prompt = prompt_template.replace("{{section}}", section)\
                           .replace("{{subsection}}", subsection)\
                           .replace("{{context}}", context)\
                           .replace("{{previous_generation}}", text)\
                           .replace("{{groundedness_score}}", str(evaluation.groundedness))\
                           .replace("{{completeness_score}}", str(evaluation.completeness))\
                           .replace("{{coherence_score}}", str(evaluation.coherence))
    return call_openai(prompt)

def generate_document_with_loop(doc_type: str, request: DocumentRequest) -> DocumentResult:
    group_id = getattr(request, 'group_id', 'default')
    sections_output = []

    for section, subsections in request.sections.items():
        context = ""
        for subsection in subsections:
            prompt_template = load_group_template(doc_type, group_id, section, subsection)
            prompt = prompt_template.replace("{{section}}", section).replace("{{subsection}}", subsection).replace("{{context}}", context)
            if (key := f"{section}::{subsection}") in request.custom_prompts:
                prompt = request.custom_prompts[key]
            generated_text = call_openai(prompt)
            evaluation = evaluate_generation(generated_text)
            # Self-reflection and improvement
            improved_text = self_reflect_and_improve(generated_text, evaluation, section, subsection, context, group_id)
            sections_output.append(SectionOutput(
                section=section,
                subsection=subsection,
                generated_text=improved_text,  # Use improved text instead of original
                evaluation=evaluation
            ))
            context += improved_text + "\n"

    return DocumentResult(document_type=doc_type, title=request.title, sections=sections_output)

def export_to_docx(document_result: DocumentResult) -> str:
    doc = Document()
    doc.add_heading(document_result.title, 0)
    current_section = None
    for section_output in document_result.sections:
        if section_output.section != current_section:
            doc.add_heading(section_output.section, level=1)
            current_section = section_output.section
        doc.add_heading(section_output.subsection, level=2)
        doc.add_paragraph(section_output.generated_text)
    # Save to a temporary file
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(tmp.name)
    return tmp.name
